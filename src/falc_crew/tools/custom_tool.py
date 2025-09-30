import os
import json
import re
from crewai.tools import BaseTool
from crewai_tools import RagTool
from typing import List, Optional, Type
from pydantic import BaseModel, Field
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ========== WordExtractorTool ==========
class WordExtractorInput(BaseModel):
    file_path: str = Field(..., description="Chemin vers le document Word source (.docx)")


class WordExtractorTool(BaseTool):
    name: str = "WordExtractorTool"
    description: str = "Lit le contenu texte d'un document Word (.docx) et retourne le texte brut."
    args_schema: Type[BaseModel] = WordExtractorInput

    def _run(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            return "⚠️ Le fichier spécifié est introuvable."

        doc = Document(file_path)
        all_paragraphs_text = []

        from docx.document import Document as _Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import _Cell, Table
        from docx.text.paragraph import Paragraph

        for block in doc.element.body:
            if isinstance(block, CT_P):
                para = Paragraph(block, doc)
                if para.text.strip():
                    all_paragraphs_text.append(para.text)
            elif isinstance(block, CT_Tbl):
                table = Table(block, doc)
                for row in table.rows:
                    for cell in row.cells:
                        for para_in_cell in cell.paragraphs:
                            if para_in_cell.text.strip():
                                all_paragraphs_text.append(para_in_cell.text)

        text = "\n".join(all_paragraphs_text)
        return text


# ========== FalcDocxStructureTaggerTool ==========

class FalcDocxStructureTaggerInput(BaseModel):
    paragraphs: List[str] = Field(..., description="List of paragraphs extracted from the original .docx file.")

class FalcDocxStructureTaggerTool(BaseTool):
    name: str = "FalcDocxStructureTaggerTool"
    description: str = (
        "Receives a list of paragraphs from the original .docx and returns the start/end "
        "indexes delimiting the section to replace with the FALC translation, "
        "and predicts the document type."
    )
    args_schema: Type[BaseModel] = FalcDocxStructureTaggerInput

    def _run(self, paragraphs: List[str]) -> str:
        # Build a numbered list for the model
        numbered = "\n".join(f"{i}. {p}" for i, p in enumerate(paragraphs))

        document_types_list = [
            "Convocation (à un rdv, à une assemblée, EDL)",
            "Décision (acceptation, refus, attribution)",
            "Courrier d'information / Notification",
            "Facture / Demande de paiement",
            "Contrat / Convention",
            "Avis (général, d'expulsion, de passage)",
            "Règlement / Procédure",
            "Demande (de documents, d'action)",
            "Attestation",
            "Autre"
        ]
        document_types_str = "\n".join(f"- {dtype}" for dtype in document_types_list)

        system_msg = {
            "role": "system",
            "content": (
                "You are an assistant that identifies the exact slice of a Word letter "
                "to replace with a FALC version, and predicts the document type. \n"
                "The 'main content' is the core message of the letter. It typically starts AFTER elements like sender's address, recipient's address, date, logos, and reference numbers. It usually begins with the subject line (if present) or the main salutation (e.g., 'Madame,', 'Objet: ...'). "
                "The 'main content' ends BEFORE final salutations (e.g., 'Veuillez agréer...'), signature blocks, or annexes not part of the core message. "
                "Header elements TO IGNORE for `replace_start` often include: postal addresses, phone numbers (Tél.), fax numbers, email addresses, website URLs, and internal references (N/réf, V/réf, Dossier). These are usually at the very top. "
                "Footer elements TO IGNORE for `replace_end` include final greetings, names and functions for signature. These are at the very bottom. "
                "Content like 'Participant details' or 'Objectif de la convention' as seen in conventions IS part of the main content. \n"
                "The main content should consist of full sentences forming a coherent message, not just isolated contact details or references. "
                "For document type prediction, consider the entire content, especially the object or main subject. \n"
                "The predicted document type MUST be one of the following: \n"
                f"{document_types_str}\n"
                "Return replace_start (index of the first paragraph of the main content), "
                "replace_end (index of the last paragraph of the main content, before any final closing/signature elements), and predicted_document_type."
            )
        }

        user_msg = {
            "role": "user",
            "content": (
                "Here are the paragraphs:\n\n"
                f"{numbered}\n\n"
                "Please analyze the content and respond *only* with valid JSON like:\n"
                '{ "replace_start": X, "replace_end": Y, "predicted_document_type": "TYPE_DE_LA_LISTE" }'
            )
        }

        from openai import OpenAI
        client = OpenAI()
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[system_msg, user_msg])

        print(f"\033[94m{response.choices[0].message.content}\033[0m")
        return response.choices[0].message.content



# ========== FalcDocxWriterTool ==========
class FalcDocxWriterInput(BaseModel):
    header: Optional[str] = Field(None, description="Header block for the letter, typically sender info")
    recipient: Optional[str] = Field(None, description="Recipient block, usually address and name")
    subject: Optional[str] = Field(None, description="Subject line of the letter")
    body_sections: List[str] = Field(..., description="List of paragraphs or markdown strings to be rendered")
    footer: Optional[str] = Field(None, description="Final line or sign-off")
    markdown_text: Optional[str] = Field(None, description="Fallback full markdown text")
    original_file: Optional[str] = Field(None, description="Path to the original .docx file")
    replace_start: int = Field(..., description="Paragraph index at which to start replacing (inclusive)")
    replace_end:   int = Field(..., description="Paragraph index at which to stop replacing (inclusive)")
    tables: Optional[List[dict]] = Field(None, description="List of table dicts with key, title, columns and rows")
    output_dir: Optional[str] = Field(None, description="Optional directory path where to save the output docx")


class FalcDocxWriterTool(BaseTool):
    name: str = "FalcDocxWriterTool"
    description: str = (
        "Generate a structured FALC Word letter layout with optional header, subject, recipient and footer"
        "This tool aslo scans text for icon placeholders (e.g. [[ICON:key]]) and inserts the corresponding PNG image."
        "It also scans for table placeholders (e.g. [[TABLE:key]]) and inserts the corresponding table."
    )
    args_schema: Type[BaseModel] = FalcDocxWriterInput

    def load_icons_map(self) -> dict:
        icons_path = os.path.abspath(os.path.join(os.getcwd(), "knowledge", "icons.json"))

        if not os.path.exists(icons_path):
            return {}
        with open(icons_path, "r", encoding="utf-8") as f:
            icon_map = json.load(f)

        # Ensure we return absolute paths
        for key, val in icon_map.items():
            if isinstance(val, str) and not os.path.isabs(val):
                # interpret val as relative to project root
                icon_map[key] = os.path.abspath(os.path.join(os.getcwd(), val))
        return icon_map

    def _clear_paragraph_formatting(self, para):
        # remove numbering and reset to Normal style
        para.style = 'Normal'
        pPr = para._p.get_or_add_pPr()
        for numPr in pPr.findall(qn('w:numPr')):
            pPr.remove(numPr)

    def _insert_text_and_icons(self, paragraph, text, icons_map):
        """
        Split the text by icon placeholders and add text runs and image runs.
        The placeholder format is assumed to be [[ICON:KEY]].
        """
        # Regex pattern to detect the placeholder pattern.
        pattern = r'\[\[ICON:(.*?)\]\]'
        # Split the text (keeping the delimiters)
        parts = re.split(f'({pattern})', text)
        for part in parts:
            if part.startswith("[[ICON:") and part.endswith("]]"):
                # Extract the key
                icon_key = re.findall(pattern, part)
                if icon_key:
                    key = icon_key[0].strip()
                    image_path = icons_map.get(key)
                    if image_path and os.path.exists(image_path):
                        run = paragraph.add_run()
                        # Adjust the image size as needed (e.g., width=Inches(0.2))
                        run.add_picture(image_path, width=Inches(0.2))
                    else:
                        # If image is not found, insert the placeholder as text.
                        paragraph.add_run(f"[Missing icon: {key}]")
            else:
                # Remove icon labels like 'direction' after [[ICON:direction]]
                cleaned = part.strip()
                if cleaned in icons_map:
                    # Skip it — already shown as an icon
                    continue

                paragraph.add_run(part)

    def _render_table_after_para(self, paragraph, table_data, icons_map, doc):
        try:
            cols = table_data.get("columns", [])
            rows = table_data.get("rows", [])

            if not cols or not rows:
                return

            body = paragraph._parent._element
            idx = list(body).index(paragraph._element)

            tmp = doc.add_table(rows=1 + len(rows), cols=len(cols))
            tmp.autofit = True

            hide_headers = table_data.get("hide_column_headers", False)

            if not hide_headers:
                for c, hdr in enumerate(cols):
                    cell = tmp.cell(0, c)
                    cell._element.clear_content()
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    self._insert_text_and_icons(cell.paragraphs[0], hdr, icons_map)

            for r, row in enumerate(rows, start=1):
                for c, val in enumerate(row):
                    cell = tmp.cell(r, c)
                    cell._element.clear_content()
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    self._insert_text_and_icons(cell.paragraphs[0], val, icons_map)

            body.insert(idx + 1, tmp._tbl)
            body.insert(idx + 2, OxmlElement("w:p"))  # spacer

        except Exception as exc:
            import logging, traceback
            logging.exception("Failed to render table")



    def _run(self, **kwargs) -> str:
        args = FalcDocxWriterInput(**kwargs)
        icons = self.load_icons_map()
        doc = Document(args.original_file)
        body = doc._body._element
        paras = list(doc.paragraphs)
        n = len(paras)
        start = max(0, min(args.replace_start, n-1))
        end   = max(0, min(args.replace_end,   n-1))
        if start > end:
            start, end = end, start

        # remove old slice
        for i in range(end, start-1, -1):
            body.remove(paras[i]._element)
        pos = start


        # insert subject
        if args.subject:
            p = doc.add_paragraph()
            self._clear_paragraph_formatting(p)
            self._insert_text_and_icons(p, args.subject, icons)
            body.insert(pos, p._p)
            pos += 1
        # insert body sections
        for sec in args.body_sections:
            sec_strip = sec.strip()
            tbl_match = re.match(r"\[\[TABLE:(.*?)\]\]", sec_strip)
            if tbl_match:
                key = tbl_match.group(1)
                table_data = {t['key']: t for t in args.tables or []}.get(key)
                placeholder = doc.add_paragraph()
                body.insert(pos, placeholder._p)

                if table_data:
                    self._render_table_after_para(placeholder, table_data, icons, doc)
                    pos += 2

                else:
                    pos += 1
            else:
                p = doc.add_paragraph()
                self._clear_paragraph_formatting(p)
                self._insert_text_and_icons(p, sec, icons)
                body.insert(pos, p._p)
                pos += 1

        # ERASE everything after the last inserted position ------------------
        trailing = list(body)[pos:]
        if trailing:
            print(f"Removing {len(trailing)} trailing paragraphs after index {pos-1}")
        for el in reversed(trailing):
            body.remove(el)

        # save
        base = os.path.splitext(os.path.basename(args.original_file))[0]
        ts = datetime.now().strftime('%Y%m%d_%H%M')
        out = args.output_dir or 'output'
        os.makedirs(out, exist_ok=True)
        dest = os.path.join(out, f"{base}_falc_{ts}.docx")
        doc.save(dest)
        return f"✅ FALC document saved: {dest}"



# ========== FalcIconLookupTool ==========
class FalcIconLookupInput(BaseModel):
    """No input needed for this tool."""
    pass

class FalcIconLookupTool(BaseTool):
    name: str = "FalcIconLookupTool"
    description: str = (
        "Provides a dictionary of available icons based on keyword-to-emoji mapping "
        "from the icons.json file, so the agent can decide where to use them."
    )
    args_schema: Type[BaseModel] = FalcIconLookupInput

    def _run(self) -> str:
        icons_path = os.path.join("knowledge", "icons.json")
        if not os.path.exists(icons_path):
            return "⚠️ Error: icons.json file not found in the 'knowledge/' directory."
        with open(icons_path, "r", encoding="utf-8") as f:
            icon_map = json.load(f)

        # Format the icons list. If the value is a PNG file, assume it needs to be rendered as an image.
        formatted_entries = []
        for key, value in icon_map.items():
            if isinstance(value, str) and value.lower().endswith(".png"):
                # Display a markdown-style image link (adjust based on your UI requirements)
                formatted_entries.append(f"- **{key}**: ![]({value})")
            else:
                formatted_entries.append(f"- **{key}**: {value}")
        formatted = "\n".join(formatted_entries)
        return f"📙 Available company icons:\n{formatted}"



# ========== ReferenceModelRetrieverTool ==========

class ReferenceModelRetrieverInput(BaseModel):
    query: str = Field(..., description="Query to search for a similar FALC translation model.")

class ReferenceModelRetrieverTool(RagTool):
    name: str = "ReferenceModelRetriever"
    description: str = (
        "Use this tool to search a bank of reference FALC translations "
        "and find similar documents based on your input."
    )
    args_schema: Type[BaseModel] = ReferenceModelRetrieverInput
