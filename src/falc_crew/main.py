#!/usr/bin/env python
import sys
import os
import base64
import warnings
import json
import openlit
import chainlit as cl
from datetime import datetime
from docx import Document
from falc_crew.crew import FalcCrew
from falc_crew.tools.custom_tool import WordExtractorTool, FalcIconLookupTool, FalcDocxStructureTaggerTool
from dotenv import load_dotenv

load_dotenv()

ENABLE_TELEMETRY = os.getenv("ENABLE_LANGFUSE_TELEMETRY", "false").lower() in ("1", "true", "yes")

if ENABLE_TELEMETRY:

    LANGFUSE_SECRET_KEY = os.getenv("LANGFUSE_SECRET_KEY")
    LANGFUSE_PUBLIC_KEY = os.getenv("LANGFUSE_PUBLIC_KEY")
    LANGFUSE_AUTH=base64.b64encode(f"{LANGFUSE_PUBLIC_KEY}:{LANGFUSE_SECRET_KEY}".encode()).decode()
    LANGFUSE_HOST = os.getenv("LANGFUSE_HOST", "https://cloud.langfuse.com")
    if not LANGFUSE_HOST.startswith(("http://", "https://")):
        LANGFUSE_HOST = "https://" + LANGFUSE_HOST

    os.environ["OTEL_EXPORTER_OTLP_ENDPOINT"] = f"{LANGFUSE_HOST}/api/public/otel"
    os.environ["OTEL_EXPORTER_OTLP_HEADERS"] = f"Authorization=Basic {LANGFUSE_AUTH}"
    os.environ["OPENAI_API_KEY"]= os.getenv("OPENAI_API_KEY")

    warnings.filterwarnings("ignore", category=SyntaxWarning, module="pysbd")

    # Initialize OpenLit for telemetry
    openlit.init()


@cl.step(name="📄 Lecture du document Word")
async def extract_text(file_path):
    extractor = WordExtractorTool()
    return extractor._run(file_path)

@cl.step(name="🔍 Analyse de la structure du document")
async def tag_structure(doc_path):
    doc = Document(doc_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tagger = FalcDocxStructureTaggerTool()
    return tagger._run(paragraphs)

@cl.step(name="🔎 Chargement des icônes disponibles")
async def load_icon_list():
    return FalcIconLookupTool()._run()

async def run(file_path: str, output_dir: str):
    print(f"📄 Lecture du fichier source : {file_path}")

    text = await extract_text(file_path)
    tag_response = await tag_structure(file_path)

    try:
        tag_data = json.loads(tag_response)
        replace_start = tag_data["replace_start"]
        replace_end = tag_data["replace_end"]
    except Exception as e:
        raise Exception(f"❌ Failed to parse structure tagging response: {tag_response}") from e

    icon_list = await load_icon_list()

    inputs = {
        "original_text": text,
        "source_filename": os.path.basename(file_path),
        "original_file": file_path,
        "replace_start": replace_start,
        "replace_end": replace_end,
        "icon_list": icon_list,
        "output_dir": output_dir,
    }

    @cl.step(name="📄 Traduction FALC en cours...")
    async def kickoff_crew(inputs):
        async with cl.Step(name="📄 Lancement", type="system") as step:
            step.input = "Texte prêt pour la traduction"
            step.output = "Analyse en cours..."

        return await FalcCrew().crew().kickoff_async(inputs=inputs)

    try:
        await kickoff_crew(inputs)
    except Exception as e:
        raise Exception(f"An error occurred while running the crew: {e}")


def train():
    """
    Train the crew using a real document from test/data/.
    Usage: uv run train <iterations> <output_filename.pkl> <docx_path (optional)>
    """
    from falc_crew.main import extract_text, tag_structure, load_icon_list

    # Prompt user for file
    train_doc_path = input("📄 Please enter the path to the .docx file you want to train on:\n> ").strip()

    if not train_doc_path or not os.path.exists(train_doc_path):
        raise Exception(f"❌ Document not found: {train_doc_path}")

    # Defaults
    iterations = 2
    # filename = f"trained_{os.path.splitext(os.path.basename(train_doc_path))[0]}.pkl"
    output_dir = "output/training"

    if not os.path.exists(train_doc_path):
        raise Exception(f"❌ Document not found: {train_doc_path}")

    def extract_text(file_path):
        extractor = WordExtractorTool()
        return extractor._run(file_path)

    def tag_structure(train_doc_path):
        doc = Document(train_doc_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tagger = FalcDocxStructureTaggerTool()
        return tagger._run(paragraphs)

    def load_icon_list():
        return FalcIconLookupTool()._run()

    # Async preprocessing
    text = extract_text(train_doc_path)
    tag_response = tag_structure(train_doc_path)

    try:
        tag_data = json.loads(tag_response)
        replace_start = tag_data["replace_start"]
        replace_end = tag_data["replace_end"]
    except Exception as e:
        raise Exception(f"❌ Structure tagging failed: {tag_response}") from e

    icon_list = load_icon_list()

    inputs = {
        "original_text": text,
        "source_filename": os.path.basename(train_doc_path),
        "original_file": train_doc_path,
        "replace_start": replace_start,
        "replace_end": replace_end,
        "icon_list": icon_list,
        "output_dir": output_dir,
    }

    # Train
    try:
        print(f"\n🏋️ Training on: {inputs['source_filename']} for {iterations} iterations")
        FalcCrew().crew().train(
            n_iterations=int(sys.argv[1]),
            filename=sys.argv[2],
            inputs=inputs
        )
        print("✅ Training complete!")
    except Exception as e:
        raise Exception(f"❌ Training failed: {e}")


def replay():
    """
    Replay the crew execution from a specific task.
    """
    try:
        FalcCrew().crew().replay(task_id=sys.argv[1])

    except Exception as e:
        raise Exception(f"An error occurred while replaying the crew: {e}")

def test():
    """
    Test the crew execution and returns the results.
    """
    inputs = {
        "topic": "AI LLMs",
        "current_year": str(datetime.now().year)
    }
    try:
        FalcCrew().crew().test(n_iterations=int(sys.argv[1]), openai_model_name=sys.argv[2], inputs=inputs)

    except Exception as e:
        raise Exception(f"An error occurred while testing the crew: {e}")


if __name__ == "__main__":
    run()

    # Inspect entity memory
    from falc_crew.crew import FalcCrew
    crew_instance = FalcCrew()
    agent = crew_instance.falc_translator()

    if agent.entity_memory:
        print("🧠 ENTITÉS MÉMORISÉES")
        for key, value in agent.entity_memory.store.items():
            print(f"{key}: {value}")
